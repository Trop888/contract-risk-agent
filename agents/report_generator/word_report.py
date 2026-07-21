"""把合同分析结果渲染成 Word 报告，返回字节流（bytes）"""
import io
from docx import Document
from docx.shared import RGBColor
from core.model import ContractAnalysisResult, RiskLevel

# 不同风险等级用不同颜色
LEVEL_COLOR = {
    RiskLevel.HIGH: RGBColor(0xC0, 0x00, 0x00),    # 红
    RiskLevel.MEDIUM: RGBColor(0xE8, 0x7D, 0x00),  # 橙
    RiskLevel.LOW: RGBColor(0x00, 0x80, 0x00),     # 绿
}


def build_word_report(result: ContractAnalysisResult) -> bytes:
    doc = Document()
    doc.add_heading("合同风险审查报告", level=0)

    # 一、总体结论
    doc.add_heading("一、总体结论", level=1)
    doc.add_paragraph(result.overall_conclusion or "—")
    p = doc.add_paragraph()
    p.add_run("整体风险等级：").bold = True
    run = p.add_run(result.overall_risk.value)
    run.bold = True
    run.font.color.rgb = LEVEL_COLOR.get(result.overall_risk)

    # 二、合同摘要
    doc.add_heading("二、合同摘要", level=1)
    doc.add_paragraph(result.summary or "—")

    # 三、风险清单
    doc.add_heading(f"三、风险清单（共 {len(result.risk_items)} 项）", level=1)
    for i, item in enumerate(result.risk_items, 1):
        doc.add_heading(f"风险 {i}：{item.risk_type}（{item.level.value}）", level=2)
        doc.add_paragraph(f"描述：{item.description}")
        if item.clause:
            doc.add_paragraph(f"相关条款：{item.clause}")
        if item.suggestion:
            doc.add_paragraph(f"修改建议：{item.suggestion}")
        if item.legal_basis:
            doc.add_paragraph(f"法律依据：{item.legal_basis}")

    # 免责声明
    doc.add_paragraph()
    disclaimer = doc.add_paragraph("免责声明：本报告由 AI 自动生成，仅供参考，不构成正式法律意见。")
    disclaimer.runs[0].italic = True

    # 存进内存字节流（不写磁盘），返回 bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()